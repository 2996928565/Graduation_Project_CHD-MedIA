"""
结构化诊断报告生成模块
基于检测结果按固定医学模板生成报告，避免依赖语言模型。
"""
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List

from loguru import logger
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _build_detection_summary(detections: List[Dict]) -> str:
    """将检测结果列表转换为自然语言摘要"""
    if not detections:
        return "未发现明显异常"

    lines = []
    for i, det in enumerate(detections, 1):
        label = det.get("label", "未知")
        conf = det.get("confidence", 0)
        measurements = det.get("measurements", {})
        bbox = det.get("bbox", [])

        line = f"{i}. 检测到：{label}（置信度：{conf:.1%}）"
        if measurements:
            w = measurements.get("width_mm", "")
            h = measurements.get("height_mm", "")
            if w and h:
                line += f"，估计大小：{w}mm × {h}mm"
        if bbox:
            line += f"，位置：({int(bbox[0])},{int(bbox[1])}) - ({int(bbox[2])},{int(bbox[3])})"
        lines.append(line)

    return "\n".join(lines)


def _build_structured_report(
    modality: str,
    patient_info: Dict,
    detections: List[Dict],
) -> Dict[str, str]:
    """根据检测结果生成固定模板结构化报告。"""
    det_summary = _build_detection_summary(detections)
    has_anomaly = any(d.get("label", "") not in ("正常",) for d in detections)

    if modality == "ultrasound":
        normal_findings = (
            "超声心动图检查：心脏各腔室大小形态未见明显异常，"
            "心肌回声尚均匀，室壁运动协调，瓣膜结构及启闭活动未见明显异常。"
            "彩色多普勒未见明显异常分流或显著反流信号。"
        )
        anomaly_findings = (
            "超声心动图检查提示心脏结构存在异常征象："
            f"{det_summary}。"
            "建议结合标准切面参数及临床表现综合评估。"
        )
        findings = anomaly_findings if has_anomaly else normal_findings
        suggestion = (
            "•提示存在先天性心脏病相关异常可能\n"
            "•建议结合心电图/心脏超声定量参数进一步评估\n"
            "•必要时行心外科或小儿心内科会诊"
            if has_anomaly
            else "•本次超声心动图检查未见明确先天性心脏病影像学异常征象"
        )
        recommendations = (
            "建议1-3个月内复查超声心动图；如有气促、发绀、喂养困难等症状，建议尽快专科就诊"
            if has_anomaly
            else "建议结合临床常规随访，必要时复查超声心动图"
        )
    else:
        normal_findings = (
            "心脏MRI（CMR）检查：心脏各腔室大小及形态未见明显异常，"
            "心肌信号分布尚均匀，室间隔及房间隔连续性可，"
            "大血管起源与走行未见明显异常，心包区未见明显积液。"
        )
        anomaly_findings = (
            "心脏MRI（CMR）检查提示结构异常征象："
            f"{det_summary}。"
            "建议结合序列参数及必要的增强检查进一步评估。"
        )
        findings = anomaly_findings if has_anomaly else normal_findings
        suggestion = (
            "•CMR提示心脏结构异常，考虑先天性心脏病相关改变可能\n"
            "•建议完善心脏功能及血流动力学相关检查\n"
            "•建议专科会诊明确分型与治疗策略"
            if has_anomaly
            else "•本次CMR未见明确心脏结构异常影像学表现"
        )
        recommendations = (
            "建议3个月内复查CMR或根据临床需要提前复查；必要时补充CTA/超声评估"
            if has_anomaly
            else "建议按临床随访计划复查，若症状变化可提前复诊"
        )

    return {
        "exam_type": "超声心动图" if modality == "ultrasound" else "心脏磁共振成像（CMR）",
        "exam_part": "心脏" if modality == "ultrasound" else "心脏及大血管",
        "image_findings": findings,
        "abnormal_findings": det_summary if has_anomaly else "未见明显异常",
        "preliminary_suggestion": suggestion,
        "recommendations": recommendations,
    }


# ──────────────────────────────────────────────────────────────────────────────
# 主报告生成接口
# ──────────────────────────────────────────────────────────────────────────────

async def generate_report(
    modality: str,
    patient_info: Dict[str, Any],
    detections: List[Dict[str, Any]],
) -> Dict[str, Any]:
    """
    生成诊断报告。

    Args:
        modality: "ultrasound" 或 "mri"
        patient_info: 患者基本信息字典
        detections: 影像检测结果列表

    Returns:
        结构化报告字典，包含 report_data（报告内容）和 metadata
    """
    patient_name = patient_info.get("name", "患者")
    patient_age = patient_info.get("age", "未知")
    patient_sex = patient_info.get("sex", "未知")
    exam_date = datetime.now().strftime("%Y年%m月%d日")

    logger.info(f"使用固定医学模板生成报告（模态: {modality}）")
    report_data = _build_structured_report(modality, patient_info, detections)
    source = "template_rule_based"

    return {
        "report_data": report_data,
        "metadata": {
            "patient_info": patient_info,
            "exam_date": exam_date,
            "modality": modality,
            "detection_count": len(detections),
            "source": source,
        },
    }


# ──────────────────────────────────────────────────────────────────────────────
# Word 导出
# ──────────────────────────────────────────────────────────────────────────────

def export_report_to_docx(report: Dict[str, Any]) -> bytes:
    """
    将报告导出为 Word (.docx) 文件字节。

    Args:
        report: generate_report() 返回的报告字典

    Returns:
        .docx 文件的字节数据
    """
    doc = Document()
    data = report.get("report_data", {})
    meta = report.get("metadata", {})
    patient_info = meta.get("patient_info", {})

    # ── 标题 ──
    title = doc.add_heading("先天性心脏病影像诊断报告（初步）", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 基本信息表格 ──
    doc.add_paragraph()
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"

    cells = info_table.rows[0].cells
    cells[0].text = "患者姓名"
    cells[1].text = str(patient_info.get("name", ""))
    cells[2].text = "年龄"
    cells[3].text = f"{patient_info.get('age', '')}岁"

    cells = info_table.rows[1].cells
    cells[0].text = "性别"
    cells[1].text = str(patient_info.get("sex", ""))
    cells[2].text = "检查日期"
    cells[3].text = str(meta.get("exam_date", ""))

    cells = info_table.rows[2].cells
    cells[0].text = "检查类型"
    cells[1].text = str(data.get("exam_type", ""))
    cells[2].text = "检查部位"
    cells[3].text = str(data.get("exam_part", ""))

    doc.add_paragraph()

    # ── 报告内容 ──
    sections = [
        ("一、影像学表现", data.get("image_findings", "")),
        ("二、异常发现", data.get("abnormal_findings", "")),
        ("三、初步诊断意见", data.get("preliminary_suggestion", "")),
        ("四、建议", data.get("recommendations", "")),
    ]

    for heading, content in sections:
        h = doc.add_heading(heading, level=2)
        para = doc.add_paragraph(content or "无")
        para.runs[0].font.size = Pt(11)

    # ── 免责声明 ──
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "【声明】本报告由系统根据检测结果按固定模板自动生成，仅供临床参考，"
        "不作为最终诊断依据。最终诊断请以临床医师结论为准。"
    )
    disclaimer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    disclaimer.runs[0].font.size = Pt(9)

    # ── 保存到字节流 ──
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
